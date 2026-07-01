"""
restorer.py — 还原核心

读取 mapping.json，将脱敏文档中的占位符逆向替换回原始实体文本。
"""

import re
import json
import logging
from pathlib import Path

import fitz
import openpyxl
import docx as python_docx

from .placeholder_registry import GlobalMapping

logger = logging.getLogger("docguard")


def load_mapping(session_dir: Path) -> GlobalMapping:
    """读取并解析 mapping.json，重建反向索引；失败时抛出 ValueError。"""
    mapping_path = session_dir / "mapping.json"
    try:
        data = json.loads(mapping_path.read_text(encoding="utf-8"))
    except Exception as e:
        raise ValueError(f"mapping.json 读取失败: {e}")
    gm = GlobalMapping(
        session_id=data["session_id"],
        created_at=data["created_at"],
        mappings=data["mappings"],
    )
    for ph, entry in data["mappings"].items():
        gm._reverse[entry["original"]] = ph
    return gm


def _build_restore_pairs(mapping: GlobalMapping) -> list[tuple[str, str]]:
    """
    构建还原替换对列表，按匹配优先级排序：
    1. 精确格式 [[P001]] → 原文（优先匹配带括号的）
    2. 宽松格式 P001 → 原文（兜底匹配大模型去掉括号的情况，用词边界确保独立 token）
    返回 [(pattern_or_string, original), ...]
    其中字符串项用 .replace()，re.Pattern 项用 .sub()
    """
    pairs: list[tuple[object, str]] = []
    # 先收集所有带括号的（按长度降序，避免 P001 被 P0011 截胡）
    bracketed = sorted(mapping.mappings.keys(), key=len, reverse=True)
    for ph in bracketed:
        original = mapping.lookup_original(ph)
        if original is not None:
            pairs.append((ph, original))
    # 再收集不带括号的宽松版本（用正则词边界）
    for ph in bracketed:
        original = mapping.lookup_original(ph)
        if original is not None:
            # 去掉 [[ ]] 得到核心标识，如 P001
            core = ph.strip("[]")
            # 只对 [A-Z]+\d{3} 格式做宽松匹配，避免误伤
            if re.fullmatch(r"[A-Z]+\d{3}", core):
                # 词边界 + 确保前后不是 [ 或 ]（避免重复匹配已带括号的）
                pattern = re.compile(r"(?<![\[\w])" + re.escape(core) + r"(?![\]\w])")
                pairs.append((pattern, original))
    return pairs


def _apply_restore(text: str, pairs: list[tuple[object, str]]) -> str:
    """对文本应用还原替换，兼容字符串和正则模式。"""
    result = text
    for pattern_or_str, original in pairs:
        if isinstance(pattern_or_str, str):
            result = result.replace(pattern_or_str, original)
        else:
            result = pattern_or_str.sub(lambda m, orig=original: orig, result)
    return result


def restore_text(text: str, mapping: GlobalMapping) -> str:
    """将文本中所有已知占位符替换回原文。兼容 [[P001]] 和 P001 两种格式。"""
    pairs = _build_restore_pairs(mapping)
    return _apply_restore(text, pairs)


def check_residual_placeholders(text: str, mapping: GlobalMapping) -> list[str]:
    """检查文本中残留的未知占位符，返回警告消息列表。兼容两种格式。"""
    warnings = []
    # 带括号格式
    pattern_bracketed = re.compile(r"\[\[[A-Z]+\d{3}\]\]")
    found_b = set(pattern_bracketed.findall(text))
    unknown_b = found_b - set(mapping.mappings.keys())
    warnings.extend(f"未知占位符保留原样: {ph}" for ph in unknown_b)
    # 不带括号格式
    pattern_loose = re.compile(r"(?<![\[\w])[A-Z]+\d{3}(?![\]\w])")
    found_l = set(pattern_loose.findall(text))
    known_cores = {ph.strip("[]") for ph in mapping.mappings.keys()}
    unknown_l = found_l - known_cores
    warnings.extend(f"疑似占位符（无括号）保留原样: {ph}" for ph in unknown_l)
    return warnings


def restore_docx(input_path: str, output_path: str, mapping: GlobalMapping) -> bool:
    """在 run 级别将占位符替换回原文，保留排版样式，保存文件。兼容两种格式。"""
    try:
        doc = python_docx.Document(input_path)
    except Exception as e:
        logger.error(f"  ❌ 读取 {input_path} 失败: {e}")
        return False

    pairs = _build_restore_pairs(mapping)

    def restore_runs(paragraphs):
        for para in paragraphs:
            for run in para.runs:
                if run.text:
                    run.text = _apply_restore(run.text, pairs)

    restore_runs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                restore_runs(cell.paragraphs)

    doc.save(output_path)
    return True


def restore_pdf(input_path: str, output_path: str, mapping: "GlobalMapping") -> bool:
    """在 PDF 中搜索占位符（兼容 [[P001]] 和 P001），用白色矩形覆盖后插入原文。"""
    try:
        doc = fitz.open(input_path)
    except Exception as e:
        logger.error(f"  ❌ 读取 {input_path} 失败: {e}")
        return False

    sorted_phs = sorted(mapping.mappings.keys(), key=len, reverse=True)

    for page in doc:
        for ph in sorted_phs:
            original = mapping.lookup_original(ph)
            if original is None:
                continue
            # 先搜带括号格式
            instances = list(page.search_for(ph))
            # 再搜不带括号的宽松格式（核心标识）
            core = ph.strip("[]")
            if re.fullmatch(r"[A-Z]+\d{3}", core):
                instances.extend(page.search_for(core))
            for inst in instances:
                page.draw_rect(inst, color=(1, 1, 1), fill=(1, 1, 1))
                fontsize = max(inst.height * 0.85, 6.0)
                try:
                    page.insert_text(inst.bl, original, fontsize=fontsize,
                                     fontname="china-s", color=(0, 0, 0))
                except Exception:
                    page.insert_text(inst.bl, original, fontsize=fontsize,
                                     color=(0, 0, 0))

    doc.save(output_path)
    doc.close()
    return True


def restore_xlsx(input_path: str, output_path: str, mapping: "GlobalMapping") -> bool:
    """遍历 Excel 所有单元格，将占位符逆向替换为原文。兼容两种格式。"""
    try:
        wb = openpyxl.load_workbook(input_path)
    except Exception as e:
        logger.error(f"  ❌ 读取 {input_path} 失败: {e}")
        return False

    pairs = _build_restore_pairs(mapping)

    for sheet in wb.worksheets:
        for row in sheet.iter_rows():
            for cell in row:
                if cell.value and isinstance(cell.value, str):
                    cell.value = _apply_restore(cell.value, pairs)

    wb.save(output_path)
    return True
