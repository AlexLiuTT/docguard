import os
import logging
import docx
from typing import List, TYPE_CHECKING
from ner_engine import extract_all_entities
from ner_engine import extract_entities_with_types
if TYPE_CHECKING:
    from reversible.placeholder_registry import GlobalMapping

logger = logging.getLogger("docguard")

def chunk_text(text: str, chunk_size: int = 1200, overlap: int = 100) -> List[str]:
    # 简单的按块划分
    chunks = []
    if len(text) <= chunk_size:
        return [text] if text else []
    
    start = 0
    while start < len(text):
        chunks.append(text[start:start+chunk_size])
        start += chunk_size - overlap
    return chunks

def process_docx(input_path: str, output_path: str):
    logger.info(f"  ▶ 正在处理 Word 文档: {os.path.basename(input_path)}")
    try:
        doc = docx.Document(input_path)
    except Exception as e:
        logger.error(f"  ❌ 读取 {input_path} 失败: {e}")
        return False, ""
    
    # 1. 提取所有纯文本，用于 NER
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        full_text.append(para.text)
                        
    raw_text = "\n".join(full_text)
    if not raw_text.strip():
        logger.warning(f"  [警告] {os.path.basename(input_path)} 未提取到文本。")
        return False, ""
    
    # 2. 分块送给大模型进行实体识别
    text_chunks = chunk_text(raw_text)
    entities = extract_all_entities(text_chunks)
    
    logger.info(f"    ✅ 共找到 {len(entities)} 个敏感实体: {entities[:5]}...")
    
    # 3. 遍历 docx 的所有 Runs，执行精准文本替换
    def replace_in_runs(paragraphs):
        for para in paragraphs:
            for run in para.runs:
                if run.text:
                    for entity in entities:
                        if entity in run.text:
                            run.text = run.text.replace(entity, "[脱敏]")

    replace_in_runs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_in_runs(cell.paragraphs)
    
    # 4. 保存文件
    doc.save(output_path)
    
    clean_text = raw_text
    for entity in entities:
        clean_text = clean_text.replace(entity, "[脱敏]")
        
    return True, clean_text


def process_docx_reversible(input_path: str, output_path: str, mapping: "GlobalMapping") -> tuple[bool, str]:
    """
    可逆脱敏版 process_docx。
    1. 提取文本 → 调用 NER（带类型）
    2. 通过 mapping.get_or_create 分配占位符
    3. run 级别替换
    4. 返回 (success, clean_text_with_placeholders)
    """
    logger.info(f"  ▶ [可逆] 正在处理 Word 文档: {os.path.basename(input_path)}")
    try:
        doc = docx.Document(input_path)
    except Exception as e:
        logger.error(f"  ❌ 读取 {input_path} 失败: {e}")
        return False, ""

    # 1. 提取所有段落和表格中的文本，拼接为 raw_text
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        full_text.append(para.text)

    raw_text = "\n".join(full_text)
    if not raw_text.strip():
        logger.warning(f"  [警告] {os.path.basename(input_path)} 未提取到文本。")
        return False, ""

    # 2. 分块送给大模型进行实体识别（带类型）
    text_chunks = chunk_text(raw_text)
    entity_pairs = extract_entities_with_types(text_chunks)  # [(original, entity_type), ...]

    logger.info(f"    ✅ 共找到 {len(entity_pairs)} 个实体（带类型）: {entity_pairs[:5]}...")

    # 3. 对每个 (original, entity_type) 调用 mapping.get_or_create 得到占位符
    sorted_entities = []
    for original, entity_type in entity_pairs:
        placeholder = mapping.get_or_create(original, entity_type)
        sorted_entities.append((original, placeholder, entity_type))

    # 按 original 长度降序，避免短实体先替换导致长实体匹配失败
    sorted_entities.sort(key=lambda x: len(x[0]), reverse=True)

    # 4. 在 run 级别替换（用占位符替代）
    def replace_in_runs(paragraphs):
        for para in paragraphs:
            for run in para.runs:
                if run.text:
                    for original, placeholder, _ in sorted_entities:
                        if original in run.text:
                            run.text = run.text.replace(original, placeholder)

    replace_in_runs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_in_runs(cell.paragraphs)

    # 5. 保存文件
    doc.save(output_path)

    # 6. 构建返回的脱敏文本
    clean_text = raw_text
    for original, placeholder, _ in sorted_entities:
        clean_text = clean_text.replace(original, placeholder)

    return True, clean_text
